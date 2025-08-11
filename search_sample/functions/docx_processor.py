# functions/docx_processor.py
"""
DOCX Processor — directly extracts paragraphs and table text from DOCX.
"""

from io import BytesIO
from docx import Document
from utils.blob_utils import download_blob_to_bytes, upload_json
from utils.db_utils import get_document_metadata
from utils.helpers import normalize_whitespace


def process_docx_direct(document_id: int, blob_path: str, output_blob_path: str):
    """
    Directly parse DOCX and upload combined text JSON.
    """
    docx_bytes = download_blob_to_bytes(blob_path)
    doc = Document(BytesIO(docx_bytes))

    text_chunks = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_chunks.append(normalize_whitespace(para.text))

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(normalize_whitespace(cell.text) for cell in row.cells)
            text_chunks.append(row_text)

    # Metadata from DB
    metadata = get_document_metadata(document_id)

    output_data = {
        "document_id": document_id,
        "metadata": metadata,
        "pages": [{"page": 1, "text": "\n".join(text_chunks)}]  # DOCX has no real pages
    }

    # Upload JSON to Azure Blob
    upload_json(output_data, output_blob_path)
    print(f"[DOCX Processor] Uploaded JSON to {output_blob_path}")
